from pathlib import Path
from copy import copy
import csv
import os
import shutil
import sys
import warnings
from datetime import datetime, timezone
from time import perf_counter
import re
import unicodedata
from collections import Counter

from cryptography.exceptions import UnsupportedAlgorithm
from cryptography.hazmat.primitives.serialization import pkcs12
from cryptography import x509
from cryptography.x509.oid import NameOID, ObjectIdentifier
from openpyxl import load_workbook
from openpyxl.styles import Alignment, PatternFill
from openpyxl.utils.cell import range_boundaries
from openpyxl.worksheet.filters import AutoFilter
from openpyxl.worksheet.table import TableColumn
from docx import Document
from dotenv import load_dotenv

from alertas import enviar_alertas, preparar_alertas, simular_alertas
from alertas_whatscontabil import (
    enviar_alertas_internos,
    preparar_alertas_internos,
    simular_alertas_internos,
)
from banco.conexao import ErroConexaoBanco, conectar_banco
from banco.consultas import (
    buscar_clientes_por_cnpjs,
    identificar_pendencias_email,
    preencher_clientes_por_nome,
    preencher_dados_clientes,
)
from integracoes.google_drive import (
    buscar_documentos_de_senha,
    conectar_google_drive,
    ler_senha_google_docs,
)
from integracoes.api_monitor import ErroIntegracaoApi, sincronizar_com_api
from utils.caminhos import obter_pasta_aplicacao
from utils.terminal import escrever, progresso, tabela, titulo


PASTA_APLICACAO = obter_pasta_aplicacao()
pasta_planilhas_locais = PASTA_APLICACAO / "planilhas"
load_dotenv(PASTA_APLICACAO / ".env")
pasta_e_cnpj = Path(os.getenv("PASTA_CERTIFICADOS", "certificados"))
arquivo_excel = Path(os.getenv("ARQUIVO_EXCEL", "relatorio_certificados.xlsx"))

# Por padrão, o terminal mostra somente o andamento e o resumo. Para uma
# investigação pontual, use TERMINAL_DETALHADO=sim no arquivo .env.
TERMINAL_DETALHADO = os.getenv("TERMINAL_DETALHADO", "nao").strip().casefold() in {
    "1", "s", "sim", "true",
}
MODO_AUTOMATICO = os.getenv("MODO_AUTOMATICO", "nao").strip().casefold() in {
    "1", "s", "sim", "true",
}
PREFIXOS_SOMENTE_DETALHADOS = (
    "=",
    "---------------------------",
    "Empresa:",
    "Certificados encontrados:",
    "Certificado:",
    "Certificado aberto com sucesso.",
    "Titular:",
    "CNPJ no arquivo:",
    "CNPJ no certificado:",
    "Vencimento:",
    "Início da validade:",
    "Senha lida",
    "Senha correspondente lida",
    "Certificado n",
    "Primeiras linhas",
    "Linha ",
    "Pend",
    "Arquivo:",
    "Status:",
    "Crit",
    "Motivo:",
    "ERRO AO ABRIR:",
    "ERRO AO LER:",
    "Detalhe técnico:",
)

# Alguns certificados antigos usam BER em vez de DER. A cryptography consegue
# fazer o fallback automaticamente; portanto este aviso não precisa poluir o
# terminal. Se o arquivo realmente falhar, o resultado continua marcado como
# ERRO NA LEITURA na planilha e no resumo.
warnings.filterwarnings(
    "ignore",
    message=r"PKCS#12 bundle could not be parsed as DER.*",
    category=UserWarning,
)


def print(*valores, **opcoes):
    """Evita mensagens repetitivas, preservando avisos, erros e resumos."""
    texto = " ".join(str(valor) for valor in valores)
    if not TERMINAL_DETALHADO and texto.startswith(PREFIXOS_SOMENTE_DETALHADOS):
        return
    if texto.startswith(("MONITOR DE CERTIFICADOS", "RELAT")) and texto.isupper():
        titulo(texto)
        return
    escrever(*valores, **opcoes)

DIAS_PRIMEIRO_AVISO = 30
DIAS_SEGUNDO_AVISO = 15
CNPJ_OID = ObjectIdentifier("2.16.76.1.3.3")
RESPONSAVEL_CERTIFICADO_OID = ObjectIdentifier("2.16.76.1.3.2")
CARACTERES_ILEGAIS_EXCEL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
INICIO_EXECUCAO = perf_counter()


def limpar_texto_excel(valor):
    """Remove caracteres de controle que uma celula do Excel nao aceita."""
    if valor is None or not isinstance(valor, str):
        return valor

    texto_limpo = CARACTERES_ILEGAIS_EXCEL.sub("", valor)
    # Uma celula do Excel aceita no maximo 32.767 caracteres.
    return texto_limpo[:32767]


def perguntar_sim_nao(pergunta, padrao=False):
    """Faz uma pergunta simples e aceita S/Sim ou N/Não."""
    if MODO_AUTOMATICO:
        if "Excel" in pergunta:
            valor = os.getenv("ATUALIZAR_EXCEL_AUTOMATICO", "sim")
        elif "sistema web" in pergunta:
            valor = os.getenv("SINCRONIZAR_API_AUTOMATICO", "nao")
        elif "WhatsCont" in pergunta:
            valor = os.getenv("ENVIAR_WHATSCONTABIL_AUTOMATICO", "nao")
        else:
            valor = os.getenv("ENVIAR_EMAIL_AUTOMATICO", "nao")
        resposta = valor.strip().casefold() in {"1", "s", "sim", "true"}
        escrever(f"Modo automático: {pergunta} {'Sim' if resposta else 'Não'}")
        return resposta

    sufixo = "[S/n]" if padrao else "[s/N]"

    while True:
        try:
            resposta = input(f"{pergunta} {sufixo}: ").strip().casefold()
        except EOFError:
            return padrao

        if not resposta:
            return padrao
        if resposta in {"s", "sim"}:
            return True
        if resposta in {"n", "nao", "não"}:
            return False

        print("Resposta inválida. Digite S para Sim ou N para Não.")


def normalizar_cnpj(valor):
    """Deixa o CNPJ no formato usado como chave nas consultas ao banco."""
    if valor is None:
        return None

    somente_numeros = re.sub(r"\D", "", str(valor))
    if len(somente_numeros) != 14:
        return None

    return somente_numeros


def indexar_certificados_por_cnpj(resultados):
    """Agrupa certificados pelo CNPJ sem perder empresas com mais de um."""
    certificados_por_cnpj = {}

    for resultado in resultados:
        cnpj = normalizar_cnpj(resultado.get("cnpj"))
        if not cnpj:
            continue

        resultado["cnpj"] = cnpj
        certificados_por_cnpj.setdefault(cnpj, []).append(resultado)

    return certificados_por_cnpj


def resumir_motivo_pendencia(motivo):
    """Transforma explicações longas em rótulos adequados ao terminal."""
    texto = str(motivo or "").casefold()
    if "cnpj não foi encontrado" in texto:
        return "CNPJ ausente no certificado e no arquivo"
    if "cadastro não localizado" in texto:
        return "Cliente não localizado no banco"
    if "e-mail não localizado" in texto:
        return "Cliente localizado, mas sem e-mail cadastrado"
    return motivo or "Motivo não informado"


def salvar_relatorios_execucao(resultados, consulta_banco_realizada):
    """Salva um resumo TXT e uma lista CSV de pendencias sem dados secretos."""
    pasta_relatorios = PASTA_APLICACAO / "relatorios"
    pasta_relatorios.mkdir(parents=True, exist_ok=True)
    identificador = datetime.now().strftime("%Y%m%d_%H%M%S")
    caminho_txt = pasta_relatorios / f"execucao_{identificador}.txt"
    caminho_csv = pasta_relatorios / f"pendencias_{identificador}.csv"
    status = Counter(item.get("status") or "SEM STATUS" for item in resultados)
    pendencias_email = (
        identificar_pendencias_email(resultados)
        if consulta_banco_realizada
        else []
    )
    erros = [
        item
        for item in resultados
        if item.get("status") not in {"EM DIA", "VENCE EM BREVE", "VENCIDO", "OK"}
    ]

    linhas = [
        "RELATORIO DA AUTOMACAO DE CERTIFICADOS",
        f"Data: {datetime.now():%d/%m/%Y %H:%M:%S}",
        f"Certificados processados: {len(resultados)}",
        f"Certificados com CNPJ: {sum(1 for item in resultados if item.get('cnpj'))}",
        f"Certificados com e-mail: {sum(1 for item in resultados if item.get('email'))}",
        f"Pendencias de e-mail: {len(pendencias_email)}",
        f"Erros de leitura: {len(erros)}",
        "",
        "QUANTIDADE POR STATUS",
    ]
    linhas.extend(f"{nome}: {quantidade}" for nome, quantidade in sorted(status.items()))
    caminho_txt.write_text("\n".join(linhas), encoding="utf-8")

    with caminho_csv.open("w", encoding="utf-8-sig", newline="") as arquivo:
        escritor = csv.writer(arquivo, delimiter=";")
        escritor.writerow(("Empresa", "CNPJ", "Arquivo", "Status", "Motivo"))
        chaves_gravadas = set()
        for item in erros:
            chave = (item.get("empresa"), item.get("arquivo"), item.get("status"))
            chaves_gravadas.add(chave)
            escritor.writerow(
                (
                    item.get("empresa") or "",
                    item.get("cnpj") or "",
                    item.get("arquivo") or "",
                    item.get("status") or "",
                    item.get("observacao") or "",
                )
            )
        for item in pendencias_email:
            chave = (item.get("empresa"), item.get("arquivo"), item.get("status"))
            if chave in chaves_gravadas:
                continue
            escritor.writerow(
                (
                    item.get("empresa") or "",
                    item.get("cnpj") or "",
                    item.get("arquivo") or "",
                    item.get("status") or "SEM E-MAIL",
                    item.get("motivo") or "",
                )
            )
    return caminho_txt, caminho_csv


def imprimir_relatorio_final(
    resultados,
    alertas,
    caminho_planilha,
    consulta_banco_realizada,
    resumo_envios,
    resumo_whatscontabil,
):
    """Exibe um resumo consolidado ao terminar toda a automação."""
    status = Counter(
        resultado.get("status") or "SEM STATUS" for resultado in resultados
    )
    origens_email = Counter(
        resultado.get("origem_email")
        for resultado in resultados
        if resultado.get("email")
    )
    pendencias = (
        identificar_pendencias_email(resultados)
        if consulta_banco_realizada
        else []
    )

    titulo("RELATÓRIO FINAL DA AUTOMAÇÃO")
    tabela(
        "Processamento",
        ("Item", "Quantidade"),
        (
            ("Certificados processados", len(resultados)),
            ("Certificados com CNPJ", sum(1 for item in resultados if item.get("cnpj"))),
            ("Certificados com e-mail", sum(1 for item in resultados if item.get("email"))),
            ("E-mails não localizados", len(pendencias)),
            ("Alertas preparados", len(alertas)),
        ),
    )
    tabela(
        "Envios",
        ("Canal", "Enviados", "Duplicados", "Falhas"),
        (
            ("E-mail", resumo_envios["enviados"], resumo_envios["duplicados"], resumo_envios["falhas"]),
            (
                "WhatsContábil",
                resumo_whatscontabil["enviados"],
                resumo_whatscontabil["duplicados"],
                resumo_whatscontabil["falhas"],
            ),
        ),
    )
    tabela(
        "Certificados por status",
        ("Status", "Quantidade"),
        tuple((nome, quantidade) for nome, quantidade in sorted(status.items())),
    )
    if origens_email:
        tabela(
            "E-mails encontrados por fonte",
            ("Fonte", "Quantidade"),
            tuple((origem, quantidade) for origem, quantidade in sorted(origens_email.items())),
        )
    if pendencias:
        tabela(
            "Empresas sem e-mail localizado",
            ("Empresa", "CNPJ", "Motivo"),
            tuple(
                (
                    item.get("empresa") or "Não informada",
                    item.get("cnpj") or "Não encontrado",
                    resumir_motivo_pendencia(item.get("motivo")),
                )
                for item in pendencias
            ),
        )
    if not consulta_banco_realizada:
        escrever("ATENÇÃO: a consulta ao banco não foi concluída.")
    escrever(f"Planilha gerada: {caminho_planilha or 'não gerada'}")
    return

    print("\n============================================================")
    print("RELATÓRIO FINAL DA AUTOMAÇÃO")
    print("============================================================")
    print(f"Certificados processados: {len(resultados)}")
    print(
        "Certificados com CNPJ: "
        f"{sum(1 for item in resultados if item.get('cnpj'))}"
    )
    print(
        "Certificados com e-mail: "
        f"{sum(1 for item in resultados if item.get('email'))}"
    )
    print(f"E-mails não localizados: {len(pendencias)}")
    print(f"Alertas de 30 ou 15 dias preparados: {len(alertas)}")
    print(f"Alertas enviados: {resumo_envios['enviados']}")
    print(f"Alertas duplicados ignorados: {resumo_envios['duplicados']}")
    print(f"Falhas de envio: {resumo_envios['falhas']}")
    print(
        "Alertas internos enviados pela WhatsContábil: "
        f"{resumo_whatscontabil['enviados']}"
    )
    print(
        "Alertas internos duplicados ignorados: "
        f"{resumo_whatscontabil['duplicados']}"
    )
    print(
        "Falhas de envio pela WhatsContábil: "
        f"{resumo_whatscontabil['falhas']}"
    )
    print("---------------------------")
    print("Quantidade por status:")
    for nome_status, quantidade in sorted(status.items()):
        print(f"{nome_status}: {quantidade}")

    if origens_email:
        print("---------------------------")
        print("E-mails encontrados por fonte:")
        for origem, quantidade in sorted(origens_email.items()):
            print(f"{origem}: {quantidade}")

    if not consulta_banco_realizada:
        print("---------------------------")
        print("A consulta de e-mails no banco não foi concluída.")
    elif pendencias:
        print("---------------------------")
        print("DETALHAMENTO DOS E-MAILS NÃO LOCALIZADOS")
        for numero, pendencia in enumerate(pendencias, start=1):
            print("---------------------------")
            print(f"Pendência {numero} de {len(pendencias)}")
            print(f"Empresa: {pendencia['empresa']}")
            print(f"CNPJ: {pendencia['cnpj']}")
            print(f"Arquivo: {pendencia['arquivo']}")
            print(f"Status: {pendencia['status']}")
            print(f"Critério: {pendencia['criterio_busca']}")
            print(f"Motivo: {pendencia['motivo']}")

    if resumo_envios["detalhes_falhas"]:
        print("---------------------------")
        print("DETALHAMENTO DAS FALHAS DE ENVIO")
        for falha in resumo_envios["detalhes_falhas"]:
            print(
                f"{falha['empresa']} | {falha['email']} | "
                f"{falha['erro']}"
            )

    print("---------------------------")
    print(f"Planilha gerada: {caminho_planilha or 'não gerada'}")
    if resumo_envios["enviados"]:
        print(
            "Mensagens aceitas pelo Gmail nesta execução: "
            f"{resumo_envios['enviados']}"
        )
    else:
        print("Nenhum e-mail de cliente foi enviado nesta execução.")
    print("============================================================")


def criar_copia_teste(arquivo_original, pasta_destino):
    if not arquivo_original.exists():
        print("Planilha original não encontrada.")
        return None

    pasta_destino.mkdir(exist_ok=True)

    arquivo_teste = pasta_destino / (
        f"{arquivo_original.stem}_TESTE{arquivo_original.suffix}"
    )

    if arquivo_teste.exists():
        print(f"Cópia de teste já existe: {arquivo_teste.name}")
        return arquivo_teste

    # copy2 cria uma cópia e preserva informações do arquivo.
    shutil.copy2(arquivo_original, arquivo_teste)
    print(f"Cópia de teste criada: {arquivo_teste.name}")

    return arquivo_teste


def mostrar_abas_da_planilha(caminho_planilha):
    if not caminho_planilha.exists():
        print("Planilha não encontrada. Verifique o caminho.")
        return

    # read_only=True abre a planilha somente para leitura.
    planilha = load_workbook(caminho_planilha, read_only=True)

    print("Planilha encontrada.")
    print(f"Abas encontradas: {planilha.sheetnames}")

    aba_certificados = planilha["Certificados"]
    print("Primeiras linhas da planilha:")
    for numero_linha, valores in enumerate(
        aba_certificados.iter_rows(
            min_row=1,
            max_row=10,
            values_only=True,
        ),
        start=1,
    ):
        print(f"Linha {numero_linha}: {valores}")

    planilha.close()


def normalizar_texto_para_busca(valor):
    texto = unicodedata.normalize("NFKD", str(valor))
    texto_sem_acentos = "".join(
        caractere
        for caractere in texto
        if not unicodedata.combining(caractere)
    )

    texto_sem_acentos = texto_sem_acentos.replace("_", " ")

    return " ".join(texto_sem_acentos.casefold().split())


def normalizar_nome_certificado_para_busca(nome_arquivo):
    caminho = Path(str(nome_arquivo))
    nome_sem_extensao = normalizar_texto_para_busca(caminho.stem)

    # Remove sufixos automáticos de cópias, sem alterar o arquivo real.
    nome_sem_extensao = re.sub(
        r"\s*\(\d+\)\s*$",
        "",
        nome_sem_extensao,
    )
    nome_sem_extensao = re.sub(
        r"^\s*copia\s+de\s+",
        "",
        nome_sem_extensao,
    )
    nome_sem_extensao = re.sub(
        r"\s*-?\s*copia(?:\s+\d+)?\s*$",
        "",
        nome_sem_extensao,
    )

    return f"{nome_sem_extensao}{caminho.suffix.casefold()}"


def normalizar_nome_certificado_exato(nome_arquivo):
    """Normaliza sem juntar o original e a cópia numerada do certificado."""
    nome = normalizar_texto_para_busca(Path(str(nome_arquivo)).name)
    # Google Drive e Excel podem divergir apenas pelo espaço antes de (1).
    return re.sub(r"\s+\((\d+)\)(?=\.[^.]+$)", r"(\1)", nome)


def atualizar_planilha(caminho_planilha, resultados):
    planilha = load_workbook(caminho_planilha)

    try:
        aba_certificados = planilha["Certificados"]

        # Usa empresa + arquivo para não confundir certificados repetidos.
        linhas_por_empresa_e_arquivo_exato = {}
        linhas_por_empresa_e_arquivo = {}
        linhas_por_arquivo_exato = {}
        linhas_por_arquivo = {}
        linhas_por_empresa = {}
        linhas_com_certificado = set()

        for numero_linha in range(8, aba_certificados.max_row + 1):
            nome_empresa = aba_certificados.cell(numero_linha, 1).value
            nome_arquivo = aba_certificados.cell(numero_linha, 2).value

            if nome_arquivo:
                linhas_com_certificado.add(numero_linha)
                arquivo_exato = normalizar_nome_certificado_exato(
                    nome_arquivo
                )
                linhas_por_arquivo_exato.setdefault(
                    arquivo_exato,
                    [],
                ).append(numero_linha)
                arquivo_normalizado = normalizar_nome_certificado_para_busca(
                    nome_arquivo
                )
                linhas_por_arquivo.setdefault(
                    arquivo_normalizado,
                    [],
                ).append(numero_linha)

                if nome_empresa:
                    empresa_normalizada = normalizar_texto_para_busca(
                        nome_empresa
                    )
                    linhas_por_empresa_e_arquivo_exato[
                        (empresa_normalizada, arquivo_exato)
                    ] = numero_linha
                    chave_completa = (
                        empresa_normalizada,
                        arquivo_normalizado,
                    )
                    linhas_por_empresa_e_arquivo[chave_completa] = (
                        numero_linha
                    )

            if nome_empresa:
                empresa_normalizada = normalizar_texto_para_busca(nome_empresa)
                linhas_por_empresa.setdefault(
                    empresa_normalizada,
                    [],
                ).append(numero_linha)

        quantidade_atualizada = 0
        quantidade_nao_encontrada = 0
        quantidade_duplicada_sem_linha = 0
        linhas_utilizadas = set()

        for resultado in resultados:
            nome_exato_para_busca = normalizar_nome_certificado_exato(
                resultado["arquivo"]
            )
            nome_para_busca = normalizar_nome_certificado_para_busca(
                resultado["arquivo"]
            )
            empresa_para_busca = normalizar_texto_para_busca(
                resultado["empresa"]
            )
            numero_linha = linhas_por_empresa_e_arquivo_exato.get(
                (empresa_para_busca, nome_exato_para_busca)
            )

            if numero_linha is None:
                numero_linha = linhas_por_empresa_e_arquivo.get(
                    (empresa_para_busca, nome_para_busca)
                )

            # Prioriza o nome exato para distinguir arquivo.pfx de
            # arquivo (1).pfx, mesmo quando houver pequena variação de espaço.
            if numero_linha is None:
                linhas_exatas = linhas_por_arquivo_exato.get(
                    nome_exato_para_busca,
                    [],
                )
                if len(linhas_exatas) == 1:
                    numero_linha = linhas_exatas[0]

            # Se o nome da empresa variar, aceita o arquivo apenas se ele for
            # único na planilha. Arquivos repetidos ficam sem atualização.
            if numero_linha is None:
                linhas_encontradas = linhas_por_arquivo.get(
                    nome_para_busca,
                    [],
                )

                if len(linhas_encontradas) == 1:
                    numero_linha = linhas_encontradas[0]

            # Último recurso: usa somente a empresa quando ela aparece uma
            # única vez. Isso cobre nomes de certificados antigos/diferentes.
            if numero_linha is None:
                linhas_da_empresa = linhas_por_empresa.get(
                    empresa_para_busca,
                    [],
                )

                if len(linhas_da_empresa) == 1:
                    numero_linha = linhas_da_empresa[0]

            if numero_linha is None:
                quantidade_nao_encontrada += 1
                continue

            if numero_linha in linhas_utilizadas:
                quantidade_duplicada_sem_linha += 1
                continue
            linhas_utilizadas.add(numero_linha)

            # C = vencimento, D = dias, E = status, F = observação.
            # Substitui sempre os valores antigos. Quando a leitura atual não
            # produz data/dias, as células ficam vazias em vez de conservar
            # informações obsoletas de uma execução anterior.
            aba_certificados.cell(numero_linha, 3).value = resultado[
                "vencimento"
            ]
            if resultado["vencimento"] is not None:
                aba_certificados.cell(numero_linha, 3).number_format = (
                    "dd/mm/yyyy"
                )

            aba_certificados.cell(numero_linha, 4).value = resultado["dias"]

            aba_certificados.cell(numero_linha, 5).value = resultado["status"]
            aba_certificados.cell(numero_linha, 6).value = limpar_texto_excel(
                resultado["observacao"]
            )

            dados_cliente = resultado.get("dados_cliente") or {}
            contatos = dados_cliente.get("contatos", [])
            socios = dados_cliente.get("socios", [])

            telefones = []
            telefone_empresa = dados_cliente.get("telefone")
            if telefone_empresa:
                telefones.append(telefone_empresa)

            for contato in contatos:
                telefone = contato.get("telefone")
                if telefone and telefone not in telefones:
                    telefones.append(telefone)

            texto_contatos = " | ".join(telefones)
            texto_socios = " | ".join(
                " - ".join(
                    parte
                    for parte in [socio.get("nome"), socio.get("email")]
                    if parte
                )
                for socio in socios
                if socio.get("nome") or socio.get("email")
            )

            aba_certificados.cell(numero_linha, 7).value = limpar_texto_excel(
                resultado.get("email") or "E-mail não localizado"
            )
            aba_certificados.cell(numero_linha, 8).value = limpar_texto_excel(
                texto_contatos or "Contato não localizado"
            )
            aba_certificados.cell(numero_linha, 9).value = limpar_texto_excel(
                texto_socios or "Nenhum sócio com e-mail localizado"
            )
            quantidade_atualizada += 1

        # Atualiza a data de geração e os indicadores do cabeçalho com as
        # linhas efetivamente processadas nesta execução.
        contagem_status = Counter(
            str(aba_certificados.cell(linha, 5).value or "").strip().upper()
            for linha in linhas_utilizadas
        )
        quantidade_vencidos = contagem_status.get("VENCIDO", 0)
        quantidade_vencendo = contagem_status.get("VENCE EM BREVE", 0)
        quantidade_em_dia = contagem_status.get("EM DIA", 0)
        quantidade_erros = sum(
            quantidade
            for status_atual, quantidade in contagem_status.items()
            if status_atual not in {"VENCIDO", "VENCE EM BREVE", "EM DIA"}
        )

        aba_certificados.cell(2, 1).value = (
            f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
        )
        aba_certificados.cell(5, 1).value = len(linhas_utilizadas)
        aba_certificados.cell(5, 2).value = quantidade_vencidos
        aba_certificados.cell(5, 3).value = quantidade_vencendo
        aba_certificados.cell(5, 4).value = quantidade_em_dia
        aba_certificados.cell(5, 5).value = quantidade_erros

        novos_cabecalhos = {
            7: "E-mail do Alerta",
            8: "Contato(s)",
            9: "Sócio(s)",
        }
        for coluna, cabecalho in novos_cabecalhos.items():
            celula = aba_certificados.cell(7, coluna, cabecalho)
            # Copia o estilo do cabeçalho existente da coluna F.
            celula._style = copy(aba_certificados.cell(7, 6)._style)

        # Reorganiza fisicamente as linhas depois de atualizar os status.
        # Antes disso, a empresa continuava na posição antiga da planilha,
        # mesmo que sua situação tivesse mudado.
        dados_tabela = [
            [
                aba_certificados.cell(numero_linha, numero_coluna).value
                for numero_coluna in range(1, 10)
            ]
            for numero_linha in range(8, aba_certificados.max_row + 1)
        ]

        prioridade_status = {
            "VENCIDO": 0,
            "VENCE EM BREVE": 1,
            "EM DIA": 2,
        }

        def chave_ordenacao(linha):
            status = str(linha[4] or "").strip().upper()
            prioridade = prioridade_status.get(status, 3)
            valor_vencimento = linha[2]

            if isinstance(valor_vencimento, datetime):
                data_ordenacao = valor_vencimento.date().isoformat()
            elif hasattr(valor_vencimento, "isoformat"):
                data_ordenacao = valor_vencimento.isoformat()
            else:
                try:
                    data_ordenacao = datetime.strptime(
                        str(valor_vencimento),
                        "%d/%m/%Y",
                    ).date().isoformat()
                except (TypeError, ValueError):
                    data_ordenacao = "9999-12-31"

            empresa = normalizar_texto_para_busca(linha[0] or "")

            # Nos grupos com data, ordena pelo vencimento. Nos erros, que
            # normalmente não possuem data, ordena pelo nome da empresa.
            if prioridade == 3:
                return (prioridade, empresa, data_ordenacao)

            return (prioridade, data_ordenacao, empresa)

        dados_tabela.sort(key=chave_ordenacao)

        for numero_linha, valores in enumerate(dados_tabela, start=8):
            for numero_coluna, valor in enumerate(valores, start=1):
                aba_certificados.cell(
                    numero_linha,
                    numero_coluna,
                ).value = valor

        # Ajusta a aparência das colunas da tabela.
        larguras = {
            "A": 42,
            "B": 65,
            "C": 16,
            "D": 18,
            "E": 20,
            "F": 55,
            "G": 36,
            "H": 25,
            "I": 50,
        }

        for letra_coluna, largura in larguras.items():
            aba_certificados.column_dimensions[letra_coluna].width = largura

        alinhamento_esquerda = Alignment(
            horizontal="left",
            vertical="center",
            wrap_text=True,
        )
        alinhamento_centro = Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True,
        )

        # As cores são recalculadas em cada execução. Isso evita que uma
        # linha com erro herde a antiga cor usada para certificados vencidos.
        preenchimentos_status = {
            "EM DIA": PatternFill("solid", fgColor="C6E0B4"),
            "VENCE EM BREVE": PatternFill("solid", fgColor="FFE699"),
            "VENCIDO": PatternFill("solid", fgColor="F4CCCC"),
        }
        preenchimento_erro = PatternFill("solid", fgColor="D9EAF7")
        preenchimento_nao_processado = PatternFill(
            "solid",
            fgColor="E7E6E6",
        )

        aba_certificados.row_dimensions[7].height = 28

        for numero_linha in range(8, aba_certificados.max_row + 1):
            aba_certificados.row_dimensions[numero_linha].height = 32

            # As colunas novas recebem bordas e fonte iguais à Observação.
            estilo_base = copy(aba_certificados.cell(numero_linha, 6)._style)
            for numero_coluna in [7, 8, 9]:
                aba_certificados.cell(numero_linha, numero_coluna)._style = (
                    copy(estilo_base)
                )

            status_linha = str(
                aba_certificados.cell(numero_linha, 5).value or ""
            ).strip().upper()

            if status_linha in preenchimentos_status:
                preenchimento_linha = preenchimentos_status[status_linha]
            elif status_linha in {"", "OK"}:
                preenchimento_linha = preenchimento_nao_processado
            else:
                # ERRO NA LEITURA, senha ausente, formato não suportado,
                # CNPJ divergente e demais situações que exigem revisão.
                preenchimento_linha = preenchimento_erro

            for numero_coluna in range(1, 10):
                aba_certificados.cell(
                    numero_linha,
                    numero_coluna,
                ).fill = preenchimento_linha

            for numero_coluna in [1, 2, 6, 7, 8, 9]:
                aba_certificados.cell(
                    numero_linha,
                    numero_coluna,
                ).alignment = alinhamento_esquerda

            for numero_coluna in [3, 4, 5]:
                aba_certificados.cell(
                    numero_linha,
                    numero_coluna,
                ).alignment = alinhamento_centro

        # Mantem sincronizados o intervalo, o autofiltro e a lista interna de
        # colunas da tabela. Alterar somente ``ref`` gera um XLSX que o Excel
        # tenta reparar ao abrir.
        for tabela in aba_certificados.tables.values():
            novo_intervalo = f"A7:I{aba_certificados.max_row}"
            cabecalhos_tabela = [
                str(aba_certificados.cell(7, coluna).value or "")
                for coluna in range(1, 10)
            ]
            colunas_anteriores = list(tabela.tableColumns)
            colunas_atualizadas = []

            for indice, cabecalho in enumerate(cabecalhos_tabela, start=1):
                if indice <= len(colunas_anteriores):
                    coluna_tabela = colunas_anteriores[indice - 1]
                    coluna_tabela.id = indice
                    coluna_tabela.name = cabecalho
                else:
                    coluna_tabela = TableColumn(
                        id=indice,
                        name=cabecalho,
                    )
                colunas_atualizadas.append(coluna_tabela)

            tabela.ref = novo_intervalo
            tabela.tableColumns = colunas_atualizadas
            tabela.autoFilter = AutoFilter(ref=novo_intervalo)

        # Impede que uma tabela inconsistente seja gravada silenciosamente.
        for tabela in aba_certificados.tables.values():
            coluna_inicial, _, coluna_final, _ = range_boundaries(tabela.ref)
            quantidade_colunas_intervalo = coluna_final - coluna_inicial + 1
            if len(tabela.tableColumns) != quantidade_colunas_intervalo:
                raise ValueError(
                    f"Tabela {tabela.name} inconsistente: o intervalo possui "
                    f"{quantidade_colunas_intervalo} colunas, mas a definição "
                    f"interna possui {len(tabela.tableColumns)}."
                )
            if tabela.autoFilter is None or tabela.autoFilter.ref != tabela.ref:
                raise ValueError(
                    f"Autofiltro da tabela {tabela.name} está inconsistente."
                )

        # Remove a antiga aba de pendências, caso tenha sido criada por uma
        # versão anterior. Os avisos agora ficam na linha de cada certificado.
        nome_aba_pendencias = "Pendências de E-mail"
        if nome_aba_pendencias in planilha.sheetnames:
            planilha.remove(planilha[nome_aba_pendencias])

        planilha.save(caminho_planilha)
        print("---------------------------")
        print(f"Linhas atualizadas no Excel: {quantidade_atualizada}")
        print(
            "Certificados sem linha correspondente: "
            f"{quantidade_nao_encontrada}"
        )
        print(
            "Certificados duplicados sem uma segunda linha no Excel: "
            f"{quantidade_duplicada_sem_linha}"
        )
        linhas_nao_associadas = linhas_com_certificado - linhas_utilizadas
        print(
            "Linhas da planilha não associadas nesta execução: "
            f"{len(linhas_nao_associadas)}"
        )
        for linha_nao_associada in sorted(linhas_nao_associadas):
            print(
                "  Linha "
                f"{linha_nao_associada}: "
                f"{aba_certificados.cell(linha_nao_associada, 1).value} | "
                f"{aba_certificados.cell(linha_nao_associada, 2).value}"
            )
    except PermissionError:
        print(
            "Não foi possível salvar a planilha. "
            "Feche o arquivo no Excel e execute novamente."
        )
    finally:
        # finally executa mesmo quando ocorre um erro.
        planilha.close()


def ler_senha(arquivo_senha):
    dados = arquivo_senha.read_bytes()

    # Tenta as codificações mais comuns encontradas no Windows.
    for codificacao in ["utf-8-sig", "utf-16", "cp1252"]:
        try:
            texto = dados.decode(codificacao)
            return texto.rstrip("\r\n")
        except UnicodeDecodeError:
            continue

    # latin-1 aceita qualquer byte e funciona como último recurso.
    return dados.decode("latin-1").rstrip("\r\n")


def ler_senha_docx(arquivo_senha):
    documento = Document(arquivo_senha)
    partes_do_texto = []

    # Lê o texto normal dos parágrafos.
    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip():
            partes_do_texto.append(paragrafo.text)

    # Lê também textos que estejam dentro de tabelas.
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if celula.text.strip():
                    partes_do_texto.append(celula.text)

    return "\n".join(partes_do_texto).rstrip("\r\n")


def criar_variacoes_da_senha(senha):
    variacoes = []

    def adicionar(valor):
        if valor and valor not in variacoes:
            variacoes.append(valor)

    # Primeiro tenta exatamente o texto encontrado no arquivo.
    adicionar(senha)

    texto_sem_invisiveis = (
        senha
        .replace("\ufeff", "")
        .replace("\u200b", "")
        .replace("\u200c", "")
        .replace("\u200d", "")
        .replace("\xa0", " ")
    )
    adicionar(texto_sem_invisiveis)

    # Depois tenta sem espaços nas pontas.
    adicionar(senha.strip())
    adicionar(texto_sem_invisiveis.strip())

    padrao_com_rotulo = re.compile(
        r"^(?:"
        r"senha(?:\s+do\s+certificado)?|"
        r"password|"
        r"documento[\s_-]*senha"
        r")"
        r"\s*(?:[:=\-\u2013\u2014]\s*|\s+)"
        r"(.+)$",
        re.IGNORECASE,
    )

    for linha in texto_sem_invisiveis.splitlines():
        linha = linha.strip()
        adicionar(linha)

        resultado_rotulo = padrao_com_rotulo.match(linha)

        if resultado_rotulo:
            adicionar(resultado_rotulo.group(1).strip())

    # Aceita arquivos que colocaram a senha entre aspas.
    for valor in variacoes.copy():
        if len(valor) >= 2 and valor[0] == valor[-1]:
            if valor[0] in ["'", '"']:
                adicionar(valor[1:-1])

    return variacoes


def abrir_certificado(arquivo_certificado, senha):
    # O PFX é um arquivo binário, por isso usamos read_bytes.
    dados_certificado = arquivo_certificado.read_bytes()

    ultimo_erro = None
    certificado_aberto = None

    for variacao in criar_variacoes_da_senha(senha):
        # A biblioteca precisa receber a senha em bytes.
        senha_em_bytes = variacao.encode("utf-8")

        try:
            chave, certificado_aberto, certificados_adicionais = (
                pkcs12.load_key_and_certificates(
                    dados_certificado,
                    senha_em_bytes,
                )
            )
            break
        except ValueError as erro:
            ultimo_erro = erro

    if certificado_aberto is None and ultimo_erro:
        raise ultimo_erro

    if certificado_aberto is None:
        raise ValueError("O PFX não possui um certificado principal.")

    return certificado_aberto


def pegar_titular(certificado_aberto):
    nomes = certificado_aberto.subject.get_attributes_for_oid(
        NameOID.COMMON_NAME
    )

    if nomes:
        return nomes[0].value

    organizacoes = certificado_aberto.subject.get_attributes_for_oid(
        NameOID.ORGANIZATION_NAME
    )

    if organizacoes:
        return organizacoes[0].value

    assunto_completo = certificado_aberto.subject.rfc4514_string()

    if assunto_completo:
        return assunto_completo

    return "Titular não encontrado"


def pegar_cnpj_do_nome(arquivo_certificado):
    resultado = re.search(
        r"(?<!\d)(\d{14})(?!\d)",
        arquivo_certificado.stem,
    )

    if resultado:
        return resultado.group(1)

    return None


def pegar_cnpj_do_certificado(certificado_aberto):
    try:
        extensao_san = certificado_aberto.extensions.get_extension_for_class(
            x509.SubjectAlternativeName
        ).value
    except x509.ExtensionNotFound:
        extensao_san = None

    if extensao_san:
        for outro_nome in extensao_san.get_values_for_type(x509.OtherName):
            if outro_nome.type_id != CNPJ_OID:
                continue

            # O valor possui marcações ASN.1 junto do texto. Alguns
            # certificados antigos também possuem bytes nulos entre dígitos.
            valor_sem_nulos = outro_nome.value.replace(b"\x00", b"")
            resultado = re.search(
                rb"(?<!\d)(\d{14})(?!\d)",
                valor_sem_nulos,
            )

            if resultado:
                return resultado.group(1).decode("ascii")

    # Alternativa encontrada em alguns certificados: serialNumber do titular.
    numeros_de_serie = certificado_aberto.subject.get_attributes_for_oid(
        NameOID.SERIAL_NUMBER
    )

    for numero_de_serie in numeros_de_serie:
        resultado = re.search(
            r"(?<!\d)(\d{14})(?!\d)",
            numero_de_serie.value,
        )

        if resultado:
            return resultado.group(1)

    # Último recurso para certificados que colocam o CNPJ no CN/assunto.
    assunto_completo = certificado_aberto.subject.rfc4514_string()
    resultado = re.search(
        r"(?<!\d)(\d{14})(?!\d)",
        assunto_completo,
    )

    if resultado:
        return resultado.group(1)

    return None


def _decodificar_texto_asn1(valor):
    """Decodifica as strings DER usadas nos campos OtherName da ICP-Brasil."""
    if not valor or len(valor) < 2:
        return None

    tag = valor[0]
    primeiro_byte_tamanho = valor[1]
    inicio = 2

    if primeiro_byte_tamanho & 0x80:
        quantidade_bytes = primeiro_byte_tamanho & 0x7F
        if not quantidade_bytes or len(valor) < inicio + quantidade_bytes:
            return None
        tamanho = int.from_bytes(
            valor[inicio:inicio + quantidade_bytes],
            "big",
        )
        inicio += quantidade_bytes
    else:
        tamanho = primeiro_byte_tamanho

    conteudo = valor[inicio:inicio + tamanho]
    if len(conteudo) != tamanho:
        return None

    # Alguns emissores envolvem outra string DER em OCTET STRING ou em um
    # campo explícito. Outros colocam o texto diretamente no OCTET STRING.
    if tag == 0x04 or tag & 0xE0 == 0xA0:
        texto_interno = _decodificar_texto_asn1(conteudo)
        if texto_interno:
            return texto_interno

        try:
            return conteudo.decode("utf-8").strip(" \t\r\n\x00") or None
        except UnicodeDecodeError:
            return None

    codificacoes = {
        0x0C: "utf-8",       # UTF8String
        0x12: "ascii",       # NumericString
        0x13: "ascii",       # PrintableString
        0x14: "latin-1",     # TeletexString
        0x16: "ascii",       # IA5String
        0x1C: "utf-32-be",   # UniversalString
        0x1E: "utf-16-be",   # BMPString
    }
    codificacao = codificacoes.get(tag)
    if not codificacao:
        return None

    try:
        return conteudo.decode(codificacao).strip(" \t\r\n\x00") or None
    except UnicodeDecodeError:
        return None


def pegar_responsavel_do_certificado(certificado_aberto):
    """Retorna o responsável de uso informado no e-CNPJ, quando disponível."""
    try:
        extensao_san = certificado_aberto.extensions.get_extension_for_class(
            x509.SubjectAlternativeName
        ).value
    except x509.ExtensionNotFound:
        return None

    for outro_nome in extensao_san.get_values_for_type(x509.OtherName):
        if outro_nome.type_id != RESPONSAVEL_CERTIFICADO_OID:
            continue

        responsavel = _decodificar_texto_asn1(outro_nome.value)
        if responsavel:
            return responsavel

    return None


def analisar_certificado(certificado, senha, resultado):
    try:
        certificado_aberto = abrir_certificado(certificado, senha)
        titular = pegar_titular(certificado_aberto)
        inicio_validade = certificado_aberto.not_valid_before_utc
        fim_validade = certificado_aberto.not_valid_after_utc
        hoje = datetime.now(timezone.utc).date()
        dias_restantes = (fim_validade.date() - hoje).days

        if dias_restantes < 0:
            status = "VENCIDO"
        elif dias_restantes <= DIAS_PRIMEIRO_AVISO:
            status = "VENCE EM BREVE"
        else:
            status = "EM DIA"

        cnpj_nome = pegar_cnpj_do_nome(certificado)
        cnpj_certificado = pegar_cnpj_do_certificado(certificado_aberto)
        responsavel_certificado = pegar_responsavel_do_certificado(
            certificado_aberto
        )

        # O CNPJ gravado dentro do certificado é a fonte principal. Se ele
        # não existir, usa o CNPJ encontrado no nome do arquivo.
        if cnpj_certificado:
            resultado["cnpj"] = normalizar_cnpj(cnpj_certificado)
            resultado["origem_cnpj"] = "certificado"
        elif cnpj_nome:
            resultado["cnpj"] = normalizar_cnpj(cnpj_nome)
            resultado["origem_cnpj"] = "nome_arquivo"

        observacoes = [f"Titular: {titular}"]

        if 0 <= dias_restantes <= DIAS_SEGUNDO_AVISO:
            observacoes.append("Faixa de aviso: 15 dias")
        elif dias_restantes <= DIAS_PRIMEIRO_AVISO and dias_restantes >= 0:
            observacoes.append("Faixa de aviso: 30 dias")

        if cnpj_nome and cnpj_certificado:
            if cnpj_nome == cnpj_certificado:
                observacoes.append(f"CNPJ conferido: {cnpj_certificado}")
            else:
                status_validade = status
                status = "CNPJ DIVERGENTE"
                observacoes.append(
                    f"CNPJ do arquivo: {cnpj_nome}; "
                    f"CNPJ do certificado: {cnpj_certificado}; "
                    f"validade: {status_validade}"
                )
        elif cnpj_certificado:
            observacoes.append(
                f"CNPJ do certificado: {cnpj_certificado}; "
                "CNPJ ausente no nome do arquivo"
            )
        elif cnpj_nome:
            observacoes.append(
                f"CNPJ do arquivo: {cnpj_nome}; "
                "CNPJ não encontrado no certificado"
            )
        else:
            observacoes.append("CNPJ não encontrado")

        resultado["vencimento"] = fim_validade.date()
        resultado["dias"] = dias_restantes
        resultado["titular"] = titular
        resultado["responsavel_certificado"] = responsavel_certificado
        resultado["status"] = status
        resultado["observacao"] = " | ".join(observacoes)

        print("Certificado aberto com sucesso.")
        print(f"Titular: {titular}")
        print(f"CNPJ no arquivo: {cnpj_nome or 'não encontrado'}")
        print(
            "CNPJ no certificado: "
            f"{cnpj_certificado or 'não encontrado'}"
        )
        print(f"Início da validade: {inicio_validade:%d/%m/%Y}")
        print(f"Vencimento: {fim_validade:%d/%m/%Y}")
    except ValueError as erro:
        resultado["observacao"] = (
            "Senha incorreta ou arquivo PFX inválido."
        )
        print("ERRO AO ABRIR: senha incorreta ou arquivo PFX inválido.")
        print(f"Detalhe técnico: {erro}")
    except UnsupportedAlgorithm as erro:
        resultado["observacao"] = "Algoritmo antigo ou não suportado."
        print("ERRO AO ABRIR: algoritmo antigo ou não suportado.")
        print(f"Detalhe técnico: {erro}")
    except OSError as erro:
        resultado["observacao"] = (
            "Erro ao ler o certificado pelo Google Drive local."
        )
        print("ERRO AO LER: certificado local indisponível.")
        print(f"Detalhe técnico: {erro}")


print("============================================================")
print("MONITOR DE CERTIFICADOS DIGITAIS")
print("============================================================")

if perguntar_sim_nao("Deseja criar ou atualizar a cópia do Excel?"):
    arquivo_excel_teste = criar_copia_teste(
        arquivo_excel,
        pasta_planilhas_locais,
    )
else:
    arquivo_excel_teste = None
    print("Cópia e atualização do Excel desativadas nesta execução.")

if arquivo_excel_teste:
    mostrar_abas_da_planilha(arquivo_excel_teste)


resultados_certificados = []
formatos_arquivos_senha = {}
drive = None
documentos_senha_por_empresa = {}

try:
    drive = conectar_google_drive(PASTA_APLICACAO)
    documentos_senha_por_empresa = buscar_documentos_de_senha(drive)
    print(
        "Google Docs de senha disponíveis: "
        f"{len(documentos_senha_por_empresa)}"
    )
except Exception as erro:
    print(
        "Não foi possível consultar os Google Docs de senha. "
        "O processamento dos arquivos locais continuará."
    )
    print(f"Detalhe técnico: {erro}")


inicio_certificados = perf_counter()
if not pasta_e_cnpj.exists():
    print("A pasta e.cnpj não existe.")
else:
    itens_drive = list(pasta_e_cnpj.iterdir())
    for pasta_empresa in progresso(
        itens_drive,
        "Analisando certificados",
        total=len(itens_drive),
    ):
        # Se o item não for uma pasta, pula para o próximo item.
        if not pasta_empresa.is_dir():
            continue

        print("---------------------------")
        print(f"Empresa: {pasta_empresa.name}")

        certificados = []
        arquivo_senha = None
        arquivo_senha_nao_suportado = None

        # Para cada item que existe dentro da pasta da empresa...
        # Algumas empresas organizam cada filial em uma subpasta. O rglob
        # percorre todos os níveis, sem deixar esses certificados de fora.
        for arquivo in pasta_empresa.rglob("*"):
            if not arquivo.is_file():
                continue

            # .suffix pega a extensão, como .pfx ou .p12.
            if arquivo.suffix.lower() in [".pfx", ".p12"]:
                certificados.append(arquivo)

            nome_sem_extensao = arquivo.stem.strip().casefold()
            extensao = arquivo.suffix.strip().casefold()

            # Também reconhece nomes como "Cópia de Senha.txt".
            if "senha" in nome_sem_extensao:
                nome_formato = extensao if extensao else "sem extensão"
                formatos_arquivos_senha[nome_formato] = (
                    formatos_arquivos_senha.get(nome_formato, 0) + 1
                )

            if (
                "senha" in nome_sem_extensao
                and extensao in [".txt", ".docx"]
            ):
                arquivo_senha = arquivo
            elif "senha" in nome_sem_extensao:
                arquivo_senha_nao_suportado = arquivo

        if certificados:
            print(f"Certificados encontrados: {len(certificados)}")
            resultados_da_empresa = [
                {
                    "empresa": pasta_empresa.name,
                    "arquivo": certificado.name,
                    # Mesmo quando o PFX não abrir, o CNPJ do nome pode ser
                    # usado para localizar o cliente e seu e-mail no banco.
                    "cnpj": normalizar_cnpj(
                        pegar_cnpj_do_nome(certificado)
                    ),
                    "origem_cnpj": (
                        "nome_arquivo"
                        if pegar_cnpj_do_nome(certificado)
                        else None
                    ),
                    # Campos reservados para o enriquecimento vindo do banco.
                    "email": None,
                    "dados_cliente": None,
                    "titular": None,
                    "responsavel_certificado": None,
                    "vencimento": None,
                    "dias": None,
                    "status": "ERRO NA LEITURA",
                    "observacao": (
                        "Não foi possível analisar o certificado."
                    ),
                }
                for certificado in certificados
            ]
        else:
            print("Certificado não encontrado.")
            resultados_da_empresa = []

        senha = None
        fonte_senha_encontrada = False

        if arquivo_senha:
            fonte_senha_encontrada = True
            erro_ao_ler_senha = False

            try:
                if arquivo_senha.suffix.casefold() == ".docx":
                    senha = ler_senha_docx(arquivo_senha)
                    print("Senha lida do arquivo Word .docx.")
                else:
                    senha = ler_senha(arquivo_senha)
                    print("Senha lida do arquivo senha.txt.")
            except Exception as erro:
                erro_ao_ler_senha = True
                print("Erro ao ler o arquivo de senha.")
                print(f"Detalhe técnico: {erro}")

                for resultado in resultados_da_empresa:
                    resultado["observacao"] = "Erro ao ler o arquivo de senha."

            if not senha and resultados_da_empresa and not erro_ao_ler_senha:
                print("O arquivo de senha está vazio.")

                for resultado in resultados_da_empresa:
                    resultado["status"] = "SENHA VAZIA"
                    resultado["observacao"] = (
                        "O arquivo de senha está vazio."
                    )
        elif drive:
            documento_senha = documentos_senha_por_empresa.get(
                pasta_empresa.name.strip().casefold()
            )

            if documento_senha:
                fonte_senha_encontrada = True
                try:
                    senha = ler_senha_google_docs(
                        drive,
                        documento_senha["id"],
                    )
                    print("Senha lida do Google Docs pela API.")
                except Exception as erro:
                    print("Erro ao ler a senha pelo Google Docs.")
                    print(f"Detalhe técnico: {erro}")

                    for resultado in resultados_da_empresa:
                        resultado["status"] = "ERRO GOOGLE DRIVE"
                        resultado["observacao"] = (
                            "Erro ao ler a senha pelo Google Docs."
                        )

        if senha and certificados:
            for certificado, resultado in zip(
                certificados,
                resultados_da_empresa,
            ):
                print(f"Certificado: {certificado.name}")

                # Dá prioridade para a senha que está na mesma subpasta
                # do certificado. Isso é necessário quando uma pasta de
                # empresa contém várias filiais, cada uma com sua senha.
                senha_certificado = senha
                arquivos_senha_da_subpasta = [
                    item
                    for item in certificado.parent.iterdir()
                    if item.is_file()
                    and "senha" in item.stem.strip().casefold()
                    and item.suffix.casefold() in [".txt", ".docx"]
                ]

                if arquivos_senha_da_subpasta:
                    arquivo_senha_especifico = arquivos_senha_da_subpasta[0]
                    try:
                        if arquivo_senha_especifico.suffix.casefold() == ".docx":
                            senha_certificado = ler_senha_docx(
                                arquivo_senha_especifico
                            )
                        else:
                            senha_certificado = ler_senha(
                                arquivo_senha_especifico
                            )
                        print(
                            "Senha correspondente lida da subpasta: "
                            f"{certificado.parent.name}"
                        )
                    except Exception as erro:
                        resultado["status"] = "ERRO NA LEITURA"
                        resultado["observacao"] = (
                            "Erro ao ler a senha da subpasta."
                        )
                        print(f"Erro ao ler senha da subpasta: {erro}")
                        continue

                analisar_certificado(
                    certificado,
                    senha_certificado,
                    resultado,
                )
        elif not senha and arquivo_senha_nao_suportado:
            extensao_senha = (
                arquivo_senha_nao_suportado.suffix.lower()
                or "sem extensão"
            )
            print(
                "Arquivo de senha encontrado, mas o formato ainda não "
                f"é suportado: {extensao_senha}"
            )
            for resultado in resultados_da_empresa:
                resultado["status"] = "FORMATO NÃO SUPORTADO"
                resultado["observacao"] = (
                    "Formato do arquivo de senha não suportado: "
                    f"{extensao_senha}"
                )
        elif not senha and not fonte_senha_encontrada:
            print("Arquivo de senha não encontrado.")

            for resultado in resultados_da_empresa:
                resultado["status"] = "SENHA NÃO ENCONTRADA"
                resultado["observacao"] = "Arquivo de senha não encontrado."

        resultados_certificados.extend(resultados_da_empresa)


print(
    "Leitura dos certificados concluída em "
    f"{perf_counter() - inicio_certificados:.1f} segundos."
)


print("---------------------------")
print("Resumo dos formatos dos arquivos de senha:")

if formatos_arquivos_senha:
    for formato, quantidade in sorted(formatos_arquivos_senha.items()):
        print(f"{formato}: {quantidade}")
else:
    print("Nenhum arquivo com nome iniciado por 'senha' foi encontrado.")


# Esta é a estrutura que será usada pela integração com o banco:
# dados_certificados_por_cnpj["12345678000199"] -> lista de certificados.
dados_certificados_por_cnpj = indexar_certificados_por_cnpj(
    resultados_certificados
)
cnpjs_para_consulta = sorted(dados_certificados_por_cnpj)

print("---------------------------")
print(f"CNPJs prontos para consulta no banco: {len(cnpjs_para_consulta)}")
print(
    "Certificados sem CNPJ para consulta: "
    f"{sum(1 for item in resultados_certificados if not item.get('cnpj'))}"
)


clientes_por_cnpj = {}
conexao_banco = None
consulta_banco_realizada = False

if cnpjs_para_consulta:
    try:
        print("---------------------------")
        print("Consultando dados dos clientes pelo CNPJ...")
        conexao_banco = conectar_banco()
        clientes_por_cnpj = buscar_clientes_por_cnpjs(
            conexao_banco,
            cnpjs_para_consulta,
        )
        preencher_dados_clientes(
            resultados_certificados,
            clientes_por_cnpj,
        )
        preencher_clientes_por_nome(
            conexao_banco,
            resultados_certificados,
        )
        consulta_banco_realizada = True

        quantidade_encontrada = sum(
            1 for cliente in clientes_por_cnpj.values() if cliente
        )
        quantidade_com_email = sum(
            1
            for resultado in resultados_certificados
            if resultado.get("email")
        )

        print(f"CNPJs encontrados no banco: {quantidade_encontrada}")
        print(
            "CNPJs não encontrados no banco: "
            f"{len(clientes_por_cnpj) - quantidade_encontrada}"
        )
        print(
            "Certificados associados a um e-mail: "
            f"{quantidade_com_email}"
        )
        print(
            "Certificados sem e-mail: "
            f"{len(resultados_certificados) - quantidade_com_email}"
        )

    except (ErroConexaoBanco, ValueError) as erro:
        print("Não foi possível consultar os clientes no banco.")
        print(f"Detalhe: {erro}")
    except Exception as erro:
        print("Erro inesperado durante a consulta dos clientes.")
        print(f"Detalhe: {erro}")
    finally:
        if conexao_banco is not None:
            conexao_banco.close()


modo_api = os.getenv("MODO_API", "desativado").strip().casefold()
if modo_api == "real" and resultados_certificados:
    if perguntar_sim_nao(
        f"Deseja sincronizar {len(resultados_certificados)} certificados com o sistema web?"
    ):
        try:
            resumo_api = sincronizar_com_api(PASTA_APLICACAO, resultados_certificados)
            print(
                "Sincronização com a API concluída: "
                f"{resumo_api.get('criados', 0)} criados, "
                f"{resumo_api.get('atualizados', 0)} atualizados e "
                f"{len(resumo_api.get('rejeitados', []))} rejeitados."
            )
        except ErroIntegracaoApi as erro:
            print(f"Falha na sincronização com a API: {erro}")


alertas_pendentes = preparar_alertas(resultados_certificados)
alertas_internos = preparar_alertas_internos(resultados_certificados)
resumo_envios = {
    "enviados": 0,
    "duplicados": 0,
    "falhas": 0,
    "detalhes_falhas": [],
    "message_ids": [],
}
resumo_whatscontabil = {
    "enviados": 0,
    "duplicados": 0,
    "falhas": 0,
    "detalhes_falhas": [],
    "message_ids": [],
}

print("---------------------------")
print(
    "Alertas de 30 ou 15 dias prontos para futura etapa de envio: "
    f"{len(alertas_pendentes)}"
)

load_dotenv(PASTA_APLICACAO / ".env")
modo_email = os.getenv("MODO_EMAIL", "simulacao").strip().casefold()

if alertas_pendentes:
    deseja_enviar = perguntar_sim_nao(
        f"Deseja enviar agora os {len(alertas_pendentes)} alertas?"
    )

    if deseja_enviar and modo_email == "real":
        print("ATENÇÃO: envio real confirmado pelo usuário.")
        resumo_envios = enviar_alertas(alertas_pendentes)
    elif deseja_enviar:
        print(
            "Envio bloqueado: configure MODO_EMAIL=real no arquivo .env."
        )
        simular_alertas(alertas_pendentes)
    else:
        print("Envio cancelado pelo usuário.")
        simular_alertas(alertas_pendentes)
else:
    simular_alertas(alertas_pendentes)


modo_whatscontabil = os.getenv(
    "MODO_WHATSCONTABIL", "desativado"
).strip().casefold()
numero_teste_whatscontabil = os.getenv(
    "WHATSCONTABIL_NUMERO_TESTE", ""
).strip()
whatsapp_id_texto = os.getenv("WHATSCONTABIL_WHATSAPP_ID", "2").strip()

print("---------------------------")
print(
    "Alertas internos entre 1 e 30 dias para a WhatsContábil: "
    f"{len(alertas_internos)}"
)

if alertas_internos:
    deseja_enviar_whatscontabil = perguntar_sim_nao(
        "Deseja testar os avisos de cliente, responsável e equipe pela "
        "WhatsContábil no número de teste?"
    )

    if deseja_enviar_whatscontabil and modo_whatscontabil == "teste":
        if not numero_teste_whatscontabil:
            print(
                "Envio bloqueado: configure "
                "WHATSCONTABIL_NUMERO_TESTE no arquivo .env."
            )
            simular_alertas_internos(alertas_internos, None)
        elif not whatsapp_id_texto.isdigit():
            print(
                "Envio bloqueado: WHATSCONTABIL_WHATSAPP_ID deve ser numérico."
            )
            simular_alertas_internos(
                alertas_internos,
                numero_teste_whatscontabil,
            )
        else:
            print(
                "ATENÇÃO: os três tipos de aviso serão enviados somente "
                "ao número de teste configurado; nenhum cliente receberá "
                "mensagem nesta fase."
            )
            resumo_whatscontabil = enviar_alertas_internos(
                alertas_internos,
                PASTA_APLICACAO,
                numero_teste_whatscontabil,
                int(whatsapp_id_texto),
            )
    elif deseja_enviar_whatscontabil:
        print(
            "Envio bloqueado: configure MODO_WHATSCONTABIL=teste no .env."
        )
        simular_alertas_internos(
            alertas_internos,
            numero_teste_whatscontabil,
        )
    else:
        print("Envio interno pela WhatsContábil cancelado pelo usuário.")
        simular_alertas_internos(
            alertas_internos,
            numero_teste_whatscontabil,
        )
else:
    simular_alertas_internos(
        alertas_internos,
        numero_teste_whatscontabil,
    )


if arquivo_excel_teste and resultados_certificados:
    atualizar_planilha(
        arquivo_excel_teste,
        resultados_certificados,
    )


imprimir_relatorio_final(
    resultados_certificados,
    alertas_pendentes,
    arquivo_excel_teste,
    consulta_banco_realizada,
    resumo_envios,
    resumo_whatscontabil,
)

try:
    relatorio_txt, relatorio_csv = salvar_relatorios_execucao(
        resultados_certificados,
        consulta_banco_realizada,
    )
    escrever(f"Resumo salvo em: {relatorio_txt}")
    escrever(f"Pendências salvas em: {relatorio_csv}")
except OSError as erro:
    escrever(f"Falha ao salvar os relatórios da execução: {erro}")

escrever(
    "Execução concluída em "
    f"{perf_counter() - INICIO_EXECUCAO:.1f} segundos."
)

if getattr(sys, "frozen", False) and not MODO_AUTOMATICO:
    input("\nPressione Enter para fechar o programa...")
